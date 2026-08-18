from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_design_spec import embed_document_font


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "研发与生产大数据平台功能手册.md"
OUT = ROOT / "docs" / "研发与生产大数据平台功能手册.docx"

DOC_FONT = "Arial Unicode MS"
DOC_FONT_FILE = Path("/Library/Fonts/Arial Unicode.ttf").resolve()

NAVY = "173B57"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "20313E"
MUTED = "657684"
SOFT = "E8EEF5"
BORDER = "CED7DF"
PALE_BLUE = "EEF5FA"
PALE_GREEN = "EAF5EF"
PALE_AMBER = "FFF6DE"
PALE_GRAY = "F6F8FA"


def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = DOC_FONT
    rpr = run._element.get_or_add_rPr()
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        rpr.rFonts.set(qn(f"w:{slot}"), DOC_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = DOC_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal._element.rPr.rFonts.set(qn(f"w:{slot}"), DOC_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = DOC_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
            style._element.rPr.rFonts.set(qn(f"w:{slot}"), DOC_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    styles = {
        "Manual Cover Title": (30, NAVY, True, WD_ALIGN_PARAGRAPH.CENTER),
        "Manual Cover Subtitle": (14, DARK_BLUE, False, WD_ALIGN_PARAGRAPH.CENTER),
        "Manual Kicker": (10, BLUE, True, WD_ALIGN_PARAGRAPH.CENTER),
        "Manual Meta": (10, INK, False, WD_ALIGN_PARAGRAPH.LEFT),
        "Table Text": (9, INK, False, WD_ALIGN_PARAGRAPH.LEFT),
        "TOC Entry": (10.5, INK, False, WD_ALIGN_PARAGRAPH.LEFT),
        "Small Note": (9, MUTED, False, WD_ALIGN_PARAGRAPH.LEFT),
    }
    for name, (size, color, bold, align) in styles.items():
        style = doc.styles[name] if name in doc.styles else doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = DOC_FONT
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
            style._element.rPr.rFonts.set(qn(f"w:{slot}"), DOC_FONT)
        style.paragraph_format.alignment = align
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.15


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def cant_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def apply_geometry(table, widths):
    assert sum(widths) == 9360, widths
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def widths_for(headers):
    count = len(headers)
    if count == 2:
        return [2300, 7060]
    if count == 3:
        return [2100, 5160, 2100]
    if count == 4:
        return [1700, 2600, 3460, 1600]
    base = 9360 // count
    return [base] * (count - 1) + [9360 - base * (count - 1)]


def table_widths(headers):
    if headers == ["编号", "导航域", "当前内容"]:
        return [900, 2400, 6060]
    if headers == ["现象/状态", "常见原因", "处理建议"]:
        return [1400, 2800, 5160]
    if len(headers) == 4 and headers[0] == "功能域":
        return [1600, 2200, 2700, 2860]
    return widths_for(headers)


def inline_runs(paragraph, text):
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_font(paragraph.add_run(text[pos:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=9.5, color=DARK_BLUE)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EEF2F5")
            run._element.get_or_add_rPr().append(shd)
        pos = match.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:]))


def add_numbering_definition(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def add_list(doc, text, num_id):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    inline_runs(p, text)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = ""
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
        set_font(p.add_run("研发与生产大数据平台｜功能手册"), size=8.5, color=MUTED)
        set_font(p.add_run("\tV1.0 · V33"), size=8.5, color=MUTED)
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_font(p.add_run("代码基线 537a0a4  ·  第 "), size=8.5, color=MUTED)
        add_field(p, "PAGE")
        set_font(p.add_run(" 页"), size=8.5, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph(style="Manual Kicker")
    p.paragraph_format.space_before = Pt(86)
    p.paragraph_format.space_after = Pt(18)
    set_font(p.add_run("JUSTEAM DATA FABRIC  ·  PRODUCT HANDBOOK"), size=10, color=BLUE, bold=True)

    p = doc.add_paragraph(style="Manual Cover Title")
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("研发与生产大数据平台"), size=30, color=NAVY, bold=True)

    p = doc.add_paragraph(style="Manual Cover Subtitle")
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run("功能手册"), size=20, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph(style="Manual Cover Subtitle")
    p.paragraph_format.space_after = Pt(42)
    set_font(p.add_run("账号安全 · 数据治理 · 数据集与证据链 · 生命周期 · 外部集成 · 系统运维"), size=12.5, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    items = [
        ("文档版本", "V1.0（当前代码 V33）"),
        ("代码基线", "master / 537a0a4"),
        ("适用对象", "业务用户、审核员、数据管理员、系统管理员与运维人员"),
        ("编制日期", "2026年8月16日"),
    ]
    for idx, (label, value) in enumerate(items):
        shade(table.rows[idx].cells[0], SOFT)
        for cell in table.rows[idx].cells:
            cell.paragraphs[0].style = "Manual Meta"
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_font(table.rows[idx].cells[0].paragraphs[0].add_run(label), size=10, color=NAVY, bold=True)
        set_font(table.rows[idx].cells[1].paragraphs[0].add_run(value), size=10, color=INK)
        cant_split(table.rows[idx])
    repeat_header(table.rows[0])
    apply_geometry(table, [2000, 7360])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE_BLUE)
    ppr.append(shd)
    set_font(p.add_run("本手册按“正式功能 / 混合展示 / 前端原型 / 外部边界”标识当前能力，避免把演示页面或待联调能力误认为已投产功能。"), size=10.5, color=DARK_BLUE, bold=True)
    doc.add_page_break()


def add_contents(doc, lines):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    inline_runs(p, "内容导航")
    p = doc.add_paragraph(style="Small Note")
    inline_runs(p, "本导航按功能域排列；各章节内继续按业务流程、权限和边界展开。")
    titles = [line[3:].strip() for line in lines if line.startswith("## ")]
    table = doc.add_table(rows=(len(titles) + 1) // 2, cols=2)
    table.style = "Table Grid"
    for idx, title in enumerate(titles):
        row = idx % ((len(titles) + 1) // 2)
        col = idx // ((len(titles) + 1) // 2)
        cell = table.rows[row].cells[col]
        cell.paragraphs[0].style = "TOC Entry"
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_font(cell.paragraphs[0].add_run(title), size=10.5, color=INK, bold=idx < 3)
        if idx % 2 == 0:
            shade(cell, PALE_GRAY)
    for row in table.rows:
        cant_split(row)
    repeat_header(table.rows[0])
    apply_geometry(table, [4680, 4680])
    doc.add_page_break()


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        parts = [part.strip() for part in lines[idx].strip().strip("|").split("|")]
        if idx == start + 1 and all(re.fullmatch(r":?-{3,}:?", part) for part in parts):
            idx += 1
            continue
        rows.append(parts)
        idx += 1
    return rows, idx


def add_table(doc, rows):
    headers = rows[0]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, SOFT)
        p = cell.paragraphs[0]
        p.style = "Table Text"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(value), size=9, color=NAVY, bold=True)
    repeat_header(table.rows[0])
    cant_split(table.rows[0])
    for row_idx, values in enumerate(rows[1:]):
        row = table.add_row()
        cant_split(row)
        if row_idx % 2 == 1:
            for cell in row.cells:
                shade(cell, "FAFBFC")
        for idx, value in enumerate(values):
            p = row.cells[idx].paragraphs[0]
            p.style = "Table Text"
            p.paragraph_format.space_after = Pt(0)
            inline_runs(p, value)
    apply_geometry(table, table_widths(headers))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_contents(doc, lines)

    bullet_id = add_numbering_definition(doc, "bullet")
    decimal_id = None
    # Let Word paginate chapters naturally. Hard page breaks left large nearly
    # empty pages whenever a preceding section happened to end near a boundary.
    break_before = set()
    first_title_skipped = False
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            decimal_id = None
            i += 1
            continue
        if line.startswith(">"):
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            decimal_id = None
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            decimal_id = None
            md_level = len(heading.group(1))
            title = heading.group(2)
            if md_level == 1 and not first_title_skipped:
                first_title_skipped = True
                i += 1
                continue
            level = max(1, md_level - 1)
            if level == 1 and title in break_before:
                doc.add_page_break()
            p = doc.add_paragraph(style=f"Heading {level}")
            inline_runs(p, title)
            i += 1
            continue
        bullet = re.match(r"^-\s+(.*)$", line)
        if bullet:
            decimal_id = None
            add_list(doc, bullet.group(1), bullet_id)
            i += 1
            continue
        number = re.match(r"^\d+\.\s+(.*)$", line)
        if number:
            if decimal_id is None:
                decimal_id = add_numbering_definition(doc, "decimal")
            add_list(doc, number.group(1), decimal_id)
            i += 1
            continue
        decimal_id = None
        p = doc.add_paragraph(style="Normal")
        inline_runs(p, line)
        i += 1

    doc.core_properties.title = "研发与生产大数据平台功能手册"
    doc.core_properties.subject = "当前代码全部功能、权限、流程与运维边界"
    doc.core_properties.author = "项目组"
    doc.core_properties.keywords = "研发数据,生产数据,数据治理,追溯,审计,生命周期,系统集成"
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    doc.settings._element.append(update)
    doc.save(OUT)
    embed_document_font(OUT, DOC_FONT_FILE, DOC_FONT)
    print(OUT)


if __name__ == "__main__":
    build()
