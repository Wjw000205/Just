from __future__ import annotations

import re
import tempfile
import uuid
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "system-design-spec.md"
OUT = ROOT / "docs" / "研发与生产大数据平台系统设计说明书.docx"
ARCH = ROOT / "docs" / "architecture-overview.png"
WORKFLOW = ROOT / "docs" / "record-workflow.png"

FONT_CN = "/System/Library/Fonts/STHeiti Medium.ttc"
FONT_CN_LIGHT = "/System/Library/Fonts/STHeiti Light.ttc"

NAVY = "123B5D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B3A"
MUTED = "667684"
LIGHT = "F2F4F7"
BORDER = "C8D1DA"
GREEN = "E7F4EC"
GREEN_TEXT = "23633B"
AMBER = "FFF4D6"
AMBER_TEXT = "7A5A00"
RED = "FDEBEC"
RED_TEXT = "9B1C1C"


def pil_font(size: int, bold: bool = False):
    return ImageFont.truetype(FONT_CN if bold else FONT_CN_LIGHT, size=size)


def rounded_box(draw, box, text, fill, outline, title_size=32, subtitle=None):
    draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = box
    font = pil_font(title_size, True)
    bbox = draw.textbbox((0, 0), text, font=font)
    tx = (x1 + x2 - (bbox[2] - bbox[0])) / 2
    ty = (y1 + y2 - (bbox[3] - bbox[1])) / 2 - (18 if subtitle else 0)
    draw.text((tx, ty), text, font=font, fill="#17324A")
    if subtitle:
        small = pil_font(20)
        bb = draw.textbbox((0, 0), subtitle, font=small)
        draw.text(((x1 + x2 - (bb[2] - bb[0])) / 2, ty + 55), subtitle, font=small, fill="#516575")


def arrow(draw, start, end, color="#5B7C99", width=5):
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    import math
    ang = math.atan2(y2 - y1, x2 - x1)
    length = 18
    for delta in (2.55, -2.55):
        p = (x2 + length * math.cos(ang + delta), y2 + length * math.sin(ang + delta))
        draw.line([end, p], fill=color, width=width)


def build_diagrams():
    img = Image.new("RGB", (1800, 930), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "研发与生产大数据平台总体架构", font=pil_font(44, True), fill="#123B5D")
    d.text((70, 104), "前后端分离 · 多存储职责分工 · PostgreSQL 权威状态", font=pil_font(24), fill="#657786")

    rounded_box(d, (70, 240, 360, 440), "Vue 3 管理端", "#E8F3FA", "#6AA6C9", 32, "动态表单 · 看板 · 追溯")
    rounded_box(d, (485, 210, 850, 470), "Spring Boot 3 API", "#EAF1FB", "#5F8FC7", 34, "认证 / 模板 / 数据集 / 设备 / 文件 / 集成")
    rounded_box(d, (1000, 180, 1260, 350), "PostgreSQL", "#E9F4EE", "#62A377", 30, "权威元数据 · 状态 · 审计")
    rounded_box(d, (1390, 180, 1650, 350), "MongoDB", "#EFF8EF", "#6AAA72", 30, "动态记录 · 状态投影")
    rounded_box(d, (1000, 455, 1260, 625), "GridFS", "#F6F1FC", "#9671B5", 30, "附件内容 · 分块存储")
    rounded_box(d, (1390, 455, 1650, 625), "Redis", "#FCEEEE", "#C77878", 30, "验证码 · 锁定 · 会话")
    rounded_box(d, (485, 665, 850, 835), "一致性与运维任务", "#FFF6E2", "#CE9A3D", 30, "Outbox · 投影修复 · 文件对账 · Actuator")

    arrow(d, (360, 340), (485, 340))
    arrow(d, (850, 300), (1000, 265))
    arrow(d, (850, 330), (1390, 265))
    arrow(d, (850, 375), (1000, 540))
    arrow(d, (850, 400), (1390, 540))
    arrow(d, (670, 470), (670, 665))
    arrow(d, (850, 750), (1000, 600), color="#B07B22")
    arrow(d, (850, 720), (1390, 335), color="#B07B22")
    img.save(ARCH, quality=96)

    img = Image.new("RGB", (1800, 620), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 35), "数据记录审核发布状态机", font=pil_font(42, True), fill="#123B5D")
    d.text((70, 88), "PostgreSQL 保存权威状态；审核与发布生成不可变电子签名", font=pil_font(23), fill="#657786")
    boxes = {
        "DRAFT": (80, 225, 330, 365),
        "PENDING_REVIEW": (455, 225, 800, 365),
        "APPROVED": (940, 145, 1210, 285),
        "PUBLISHED": (1410, 145, 1700, 285),
        "REJECTED": (940, 385, 1210, 525),
    }
    labels = {
        "DRAFT": ("DRAFT", "草稿"),
        "PENDING_REVIEW": ("PENDING_REVIEW", "待独立审核"),
        "APPROVED": ("APPROVED", "审核通过"),
        "PUBLISHED": ("PUBLISHED", "已发布"),
        "REJECTED": ("REJECTED", "已驳回"),
    }
    fills = {"DRAFT": "#EAF1FB", "PENDING_REVIEW": "#FFF4D6", "APPROVED": "#E7F4EC", "PUBLISHED": "#DFF3EA", "REJECTED": "#FDEBEC"}
    outlines = {"DRAFT": "#5F8FC7", "PENDING_REVIEW": "#CE9A3D", "APPROVED": "#62A377", "PUBLISHED": "#3F9363", "REJECTED": "#C77878"}
    for key, box in boxes.items():
        rounded_box(d, box, labels[key][0], fills[key], outlines[key], 27, labels[key][1])
    arrow(d, (330, 295), (455, 295))
    d.text((345, 252), "提交", font=pil_font(19), fill="#516575")
    arrow(d, (800, 270), (940, 215))
    d.text((835, 205), "批准+签名", font=pil_font(18), fill="#516575")
    arrow(d, (1210, 215), (1410, 215))
    d.text((1240, 172), "再认证发布", font=pil_font(18), fill="#516575")
    arrow(d, (800, 330), (940, 450))
    d.text((825, 390), "驳回+签名", font=pil_font(18), fill="#516575")
    arrow(d, (940, 500), (330, 365), color="#A66C45")
    d.text((525, 456), "受控更正并升版后回到草稿", font=pil_font(18), fill="#79523B")
    d.rounded_rectangle((1280, 390, 1710, 525), radius=16, fill="#F7F9FB", outline="#AAB6C0", width=2)
    d.text((1310, 415), "职责分离规则", font=pil_font(23, True), fill="#294B65")
    d.text((1310, 458), "创建者不可审核自己的记录", font=pil_font(19), fill="#516575")
    d.text((1310, 490), "通过/发布后禁止原地修改", font=pil_font(19), fill="#516575")
    img.save(WORKFLOW, quality=96)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def apply_table_geometry(table, widths):
    assert sum(widths) == 9360, widths
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "9360")
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[idx]))
            tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# LibreOffice headless may expose the macOS Hiragino collection under an
# internal dot-prefixed family and render CJK glyphs as empty boxes.  Arial
# Unicode MS is present in the supported desktop/runtime image and keeps the
# generated DOCX readable in both Word and the required headless PDF preview.
# STSong is available to both macOS Word and the standard LibreOffice
# headless renderer, avoiding environment-specific CJK fallback failures.
DOC_FONT = "Arial Unicode MS"
DOC_FONT_FILE = Path("/Library/Fonts/Arial Unicode.ttf").resolve()


def set_font(run, size=None, color=None, bold=None, italic=None, code=False):
    # LibreOffice on macOS may ignore w:eastAsia when w:ascii/w:hAnsi point to
    # a Latin-only family.  Use one installed CJK-capable family for every font
    # slot so Chinese text renders consistently in Word and headless previews.
    name = DOC_FONT
    east = DOC_FONT
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = DOC_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = DOC_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name, size, color, align in (
        ("Spec Title", 28, NAVY, WD_ALIGN_PARAGRAPH.LEFT),
        ("Spec Subtitle", 14, MUTED, WD_ALIGN_PARAGRAPH.LEFT),
        ("Table Text", 9, INK, WD_ALIGN_PARAGRAPH.LEFT),
        ("Figure Caption", 9, MUTED, WD_ALIGN_PARAGRAPH.CENTER),
        ("TOC Entry", 11, INK, WD_ALIGN_PARAGRAPH.LEFT),
    ):
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = DOC_FONT
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
        style.paragraph_format.alignment = align
        style.paragraph_format.space_after = Pt(4 if name != "Spec Title" else 8)
        style.paragraph_format.line_spacing = 1.1


def add_numbering_definition(doc, kind):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing_abs, default=0) + 1
    existing_num = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    num_id = max(existing_num, default=0) + 1
    marker = "•" if kind == "bullet" else "%1."
    fmt = "bullet" if kind == "bullet" else "decimal"
    abstract = parse_xml(
        f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{abstract_id}">'
        '<w:multiLevelType w:val="singleLevel"/>'
        '<w:lvl w:ilvl="0">'
        '<w:start w:val="1"/>'
        f'<w:numFmt w:val="{fmt}"/>'
        f'<w:lvlText w:val="{marker}"/>'
        '<w:lvlJc w:val="left"/>'
        '<w:pPr><w:tabs><w:tab w:val="num" w:pos="720"/></w:tabs><w:ind w:left="720" w:hanging="360"/></w:pPr>'
        f'<w:rPr><w:rFonts w:ascii="{DOC_FONT}" w:hAnsi="{DOC_FONT}" w:eastAsia="{DOC_FONT}"/><w:sz w:val="22"/></w:rPr>'
        '</w:lvl></w:abstractNum>'
    )
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def add_list_paragraph(doc, text, num_id):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    numPr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    add_inline_runs(p, text)
    return p


def add_inline_runs(paragraph, text):
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_font(paragraph.add_run(text[pos:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=9.5, color=DARK_BLUE, code=True)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EEF2F5")
            run._element.get_or_add_rPr().append(shd)
        pos = match.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:]))


def widths_for(headers):
    n = len(headers)
    if n == 2:
        return [2300, 7060]
    if n == 3:
        return [2100, 5160, 2100]
    if n == 4:
        if any("主要表" in h for h in headers):
            return [1500, 4100, 2600, 1160]
        return [1700, 2350, 3710, 1600]
    return [9360 // n] * (n - 1) + [9360 - (9360 // n) * (n - 1)]


def add_table(doc, rows):
    headers = rows[0]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.style = "Table Text"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=9, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    for row_values in rows[1:]:
        row = table.add_row()
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        cells = row.cells
        for i, value in enumerate(row_values):
            cell = cells[i]
            p = cell.paragraphs[0]
            p.style = "Table Text"
            p.paragraph_format.space_after = Pt(0)
            if len(value) <= 14 and i in (0, len(row_values) - 1):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, value)
            if "已实现并测试" in value or "已实现并验证" in value or value == "已实现并测试":
                set_cell_shading(cell, GREEN)
                for run in p.runs:
                    run.font.color.rgb = RGBColor.from_string(GREEN_TEXT)
            elif "基础能力" in value or "待压测" in value or "待专项验收" in value or "待版本矩阵" in value or "待验收" in value:
                set_cell_shading(cell, AMBER)
                for run in p.runs:
                    run.font.color.rgb = RGBColor.from_string(AMBER_TEXT)
            elif "待定/未纳入" in value or "外部阻塞" in value:
                set_cell_shading(cell, RED)
                for run in p.runs:
                    run.font.color.rgb = RGBColor.from_string(RED_TEXT)
    apply_table_geometry(table, widths_for(headers))
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    return table


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = ""
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
        r = p.add_run("研发与生产大数据平台｜系统设计说明书")
        set_font(r, size=8.5, color=MUTED)
        r = p.add_run("\tV1.9")
        set_font(r, size=8.5, color=MUTED)

        footer = section.footer
        p = footer.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run("第 ")
        set_font(r, size=8.5, color=MUTED)
        add_field(p, "PAGE")
        r = p.add_run(" 页")
        set_font(r, size=8.5, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("研发与生产大数据平台数据库建设")
    set_font(r, size=12, color=BLUE, bold=True)

    p = doc.add_paragraph(style="Spec Title")
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("系统设计说明书")
    set_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph(style="Spec Subtitle")
    p.paragraph_format.space_after = Pt(32)
    r = p.add_run("架构设计 · 详细设计 · 数据库设计｜实现对齐版")
    set_font(r, size=14, color=MUTED)

    for label, value in (
        ("文档版本", "V1.9（V33 + BAK-002 跨库恢复对齐版）"),
        ("技术基线", "Spring Boot 3.5 / Vue 3 / PostgreSQL 17 / MongoDB 8 / Redis 8"),
        ("对齐依据", "软件开发需求文档-0408、项目实施方案修改2、需求基线与当前代码"),
        ("编制日期", "2026年8月12日"),
        ("文档状态", "核心开发/演示设计基线；生产门禁另行验收"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"{label}：")
        set_font(r, size=10.5, color=NAVY, bold=True)
        r = p.add_run(value)
        set_font(r, size=10.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EEF4F8")
    pPr.append(shd)
    r = p.add_run("设计口径：PostgreSQL 是权限、业务元数据、工作流、电子签名与审计的权威来源；MongoDB/GridFS 承载动态记录与附件；Redis 承载有期限的安全状态。")
    set_font(r, size=10.5, color=DARK_BLUE, bold=True)
    doc.add_page_break()


def add_contents(doc, markdown_lines):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.add_run("目录")
    for line in markdown_lines:
        if re.match(r"^# ", line):
            title = line[2:].strip()
            p = doc.add_paragraph(style="TOC Entry")
            p.paragraph_format.left_indent = Inches(0)
            r = p.add_run(title)
            set_font(r, size=11, color=INK, bold=True)
        elif re.match(r"^## ", line):
            title = line[3:].strip()
            p = doc.add_paragraph(style="TOC Entry")
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(title)
            set_font(r, size=9.5, color=MUTED)
    doc.add_page_break()


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        parts = [p.strip() for p in lines[idx].strip().strip("|").split("|")]
        if idx == start + 1 and all(re.fullmatch(r":?-{3,}:?", p) for p in parts):
            idx += 1
            continue
        rows.append(parts)
        idx += 1
    return rows, idx


def build_docx():
    build_diagrams()
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_header_footer(doc)
    bullet_id = add_numbering_definition(doc, "bullet")
    add_cover(doc)
    add_contents(doc, lines)

    i = 0
    decimal_id = None
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            decimal_id = None
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            decimal_id = None
            rows, i = parse_table(lines, i)
            if rows:
                add_table(doc, rows)
            continue
        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            decimal_id = None
            caption, rel = image_match.groups()
            path = SOURCE.parent / rel
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            shape = p.add_run().add_picture(str(path), width=Inches(6.2))
            shape._inline.docPr.set("descr", caption)
            shape._inline.docPr.set("title", caption)
            cap = doc.add_paragraph(style="Figure Caption")
            cap.paragraph_format.keep_with_next = True
            r = cap.add_run(caption)
            set_font(r, size=9, color=MUTED)
            i += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            decimal_id = None
            level = len(heading.group(1))
            title = heading.group(2)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline_runs(p, title)
            i += 1
            continue
        bullet = re.match(r"^-\s+(.*)$", line)
        if bullet:
            decimal_id = None
            add_list_paragraph(doc, bullet.group(1), bullet_id)
            i += 1
            continue
        number = re.match(r"^\d+\.\s+(.*)$", line)
        if number:
            if decimal_id is None:
                decimal_id = add_numbering_definition(doc, "decimal")
            add_list_paragraph(doc, number.group(1), decimal_id)
            i += 1
            continue
        decimal_id = None
        p = doc.add_paragraph(style="Normal")
        add_inline_runs(p, line)
        i += 1

    doc.core_properties.title = "研发与生产大数据平台系统设计说明书"
    doc.core_properties.subject = "架构设计、详细设计、数据库设计（实现对齐版）"
    doc.core_properties.author = "项目组"
    doc.core_properties.keywords = "Spring Boot, Vue 3, PostgreSQL, MongoDB, Redis, 系统设计"
    settings = doc.settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)
    doc.save(OUT)
    embed_document_font(OUT, DOC_FONT_FILE, DOC_FONT)
    print(OUT)


def embed_document_font(docx_path: Path, font_path: Path, family_name: str):
    """Embed the CJK body font so headless and non-macOS renderers are stable.

    The macOS system font is licensed for editable embedding (OS/2 fsType=0x8).
    OOXML stores an obfuscated copy whose first 32 bytes are XORed with the
    reversed font-key bytes.  LibreOffice and Word can then render Chinese
    without relying on a machine-specific fontconfig setup.
    """
    if not font_path.is_file():
        raise FileNotFoundError(f"Document font not found: {font_path}")

    font_key = uuid.uuid4()
    key = font_key.bytes_le
    payload = bytearray(font_path.read_bytes())
    for index in range(min(32, len(payload))):
        payload[index] ^= key[15 - (index % 16)]

    with zipfile.ZipFile(docx_path, "r") as source:
        parts = {info.filename: source.read(info.filename) for info in source.infolist()}

    content_types = etree.fromstring(parts["[Content_Types].xml"])
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    if not content_types.xpath("ct:Default[@Extension='odttf']", namespaces={"ct": ct_ns}):
        default = etree.SubElement(content_types, f"{{{ct_ns}}}Default")
        default.set("Extension", "odttf")
        default.set("ContentType", "application/vnd.openxmlformats-officedocument.obfuscatedFont")
    parts["[Content_Types].xml"] = etree.tostring(content_types, xml_declaration=True, encoding="UTF-8", standalone=True)

    rels_path = "word/_rels/fontTable.xml.rels"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    if rels_path in parts:
        rels = etree.fromstring(parts[rels_path])
    else:
        rels = etree.Element(f"{{{rel_ns}}}Relationships", nsmap={None: rel_ns})
    existing_ids = {node.get("Id") for node in rels}
    rel_index = 1
    while f"rId{rel_index}" in existing_ids:
        rel_index += 1
    rel_id = f"rId{rel_index}"
    relationship = etree.SubElement(rels, f"{{{rel_ns}}}Relationship")
    relationship.set("Id", rel_id)
    relationship.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font")
    relationship.set("Target", "fonts/ArialUnicodeMS.odttf")
    parts[rels_path] = etree.tostring(rels, xml_declaration=True, encoding="UTF-8", standalone=True)

    font_table = etree.fromstring(parts["word/fontTable.xml"])
    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    office_rel_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    for existing in font_table.xpath("w:font[@w:name=$name]", namespaces={"w": word_ns}, name=family_name):
        font_table.remove(existing)
    font = etree.SubElement(font_table, f"{{{word_ns}}}font")
    font.set(f"{{{word_ns}}}name", family_name)
    embed = etree.SubElement(font, f"{{{word_ns}}}embedRegular")
    embed.set(f"{{{office_rel_ns}}}id", rel_id)
    embed.set(f"{{{word_ns}}}fontKey", "{" + str(font_key).upper() + "}")
    parts["word/fontTable.xml"] = etree.tostring(font_table, xml_declaration=True, encoding="UTF-8", standalone=True)
    parts["word/fonts/ArialUnicodeMS.odttf"] = bytes(payload)

    with tempfile.NamedTemporaryFile(prefix="rdp-design-", suffix=".docx", dir=docx_path.parent, delete=False) as handle:
        replacement = Path(handle.name)
    try:
        with zipfile.ZipFile(replacement, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as target:
            for name, data in parts.items():
                target.writestr(name, data)
        replacement.replace(docx_path)
    finally:
        replacement.unlink(missing_ok=True)


if __name__ == "__main__":
    build_docx()
